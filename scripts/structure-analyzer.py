import os
import subprocess
import sys

# Required packages and their pip names
required_packages = {
    "pandas": "pandas",
    "docx": "python-docx",
    "fpdf": "fpdf"
}

def check_pip_available():
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "--version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except:
        return False

def install_missing_packages():
    if not check_pip_available():
        print("❌ pip is not installed. Please install pip first using:")
        print("👉 Download get-pip.py from: https://bootstrap.pypa.io/get-pip.py")
        print("👉 Then run: python get-pip.py")
        sys.exit(1)

    for module_name, package_name in required_packages.items():
        try:
            __import__(module_name)
        except ImportError:
            print(f"📦 Installing {package_name}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

# Install missing dependencies if needed
install_missing_packages()

# Now safely import packages
import pandas as pd
from docx import Document
from fpdf import FPDF

# Get folder structure
def get_folder_structure(folder_path):
    structure = []
    for root, dirs, files in os.walk(folder_path):
        level = root.replace(folder_path, '').count(os.sep)
        indent = ' ' * 4 * level
        structure.append(f"{indent}{os.path.basename(root)}/")
        for f in files:
            structure.append(f"{indent}{' ' * 4}{f}")
    return structure

# Export functions
def export_to_excel(structure, filename):
    df = pd.DataFrame(structure, columns=["Folder Structure"])
    df.to_excel(filename, index=False)
    print(f"✅ Exported to Excel: {filename}")

def export_to_word(structure, filename):
    doc = Document()
    doc.add_heading('Folder Structure', level=1)
    for line in structure:
        doc.add_paragraph(line)
    doc.save(filename)
    print(f"✅ Exported to Word: {filename}")

def export_to_pdf(structure, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Folder Structure", ln=True, align="C")
    for line in structure:
        pdf.cell(200, 8, txt=line, ln=True)
    pdf.output(filename)
    print(f"✅ Exported to PDF: {filename}")

# User interaction
def ask_user_for_location():
    print("📂 Folder Structure Export Tool\n")
    print("Do you want to analyze the current location?")
    print("1. Yes (use current directory)")
    print("2. No (enter custom directory path)")
    
    choice = input("Enter 1 or 2: ").strip()
    if choice == '1':
        folder_path = os.getcwd()
    elif choice == '2':
        folder_path = input("Paste the full custom directory path: ").strip()
        if not os.path.exists(folder_path):
            print("❌ Error: The provided path does not exist.")
            return
    else:
        print("❌ Invalid choice. Please enter 1 or 2.")
        return

    print(f"\n📁 Analyzing folder: {folder_path}\n")
    structure = get_folder_structure(folder_path)
    
    # Preview folder structure
    for line in structure:
        print(line)

    # Export options
    print("\n📤 Select export format:")
    print("1. Excel (.xlsx)")
    print("2. Word (.docx)")
    print("3. PDF (.pdf)")
    export_choice = input("Enter 1, 2, or 3: ").strip()
    
    export_name = input("Enter filename (without extension): ").strip()
    
    if export_choice == '1':
        export_to_excel(structure, f"{export_name}.xlsx")
    elif export_choice == '2':
        export_to_word(structure, f"{export_name}.docx")
    elif export_choice == '3':
        export_to_pdf(structure, f"{export_name}.pdf")
    else:
        print("❌ Invalid choice.")

# Run it
ask_user_for_location()
